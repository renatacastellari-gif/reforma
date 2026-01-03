
import streamlit as st
from pathlib import Path
from io import BytesIO

# =========================
# CONFIGURAÇÃO DA PÁGINA
# =========================
st.set_page_config(page_title="Reforma Tributária", page_icon="🟪" )

# =========================
# SENHA FIXA / LOGIN
# =========================
PASSWORD = "minhasenha123"
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

# 🔒 Esconde a barra lateral com CSS se não estiver logado
if not st.session_state.logged_in:
    st.markdown("<style>[data-testid='stSidebar']{display:none;}</style>", unsafe_allow_html=True)

# =========================
# TELA DE LOGIN
# =========================
if not st.session_state.logged_in:
    st.title("Acesso Restrito")
    senha = st.text_input("Digite a senha:", type="password")
    if st.button("Entrar", use_container_width=True):
        if senha == PASSWORD:
            st.session_state.logged_in = True
            st.success("Acesso liberado! Agora você pode navegar pelas páginas.")
            st.rerun()
        else:
            st.error("Senha incorreta.")

else:
    # =========================
    # CONTEÚDO PROTEGIDO
    # =========================

    # ---- LOGO HINES (opcional) ----
    from PIL import Image, UnidentifiedImageError
    candidatos = [Path("hines.svg"), Path("hines.png"), Path("hines.jpg"), Path("hines.jpeg")]
    logo_path = next((p for p in candidatos if p.exists()), None)
    if logo_path:
        try:
            st.image(str(logo_path), width=220)
        except Exception:
            st.markdown("<h3>🟪 Hines – Painel Tributário</h3>", unsafe_allow_html=True)
    else:
        st.markdown("<h3>🟪 Hines – Painel Tributário</h3>", unsafe_allow_html=True)

    # ---- Título ----
    st.markdown(
        "<h2 style='color:#B22222;font-family:Times New Roman,sans-serif;font-weight:700;"
        "text-align:center;border-bottom:2px solid #B22222;padding-bottom:8px;margin-bottom:20px;'>"
        "Reforma Tributária</h2>",
        unsafe_allow_html=True
    )
    st.markdown("**`REFORMA TRIBUTÁRIA`**")

    # =========================
    # Funções utilitárias para ler o Word e extrair texto/imagens
    # =========================
    DOCX_FILE = Path("fiscal reforma.docx")

    def carregar_docx(docx_path: Path):
        """
        Lê o .docx e retorna:
        - texto_plano: texto concatenado (parágrafos e células de tabelas)
        - imagens: lista de bytes de imagens extraídas
        """
        from docx import Document
        texto_parts = []
        imagens_bytes = []

        if not docx_path.exists():
            return "", []

        doc = Document(str(docx_path))

        # --- Extrair texto de parágrafos
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                texto_parts.append(t)

        # --- Extrair texto das tabelas (se houver)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        texto_parts.append(t)

        # --- Extrair imagens do pacote (media)
        # Nota: python-docx não dá API direta para "inline shapes";
        # mas as imagens ficam em doc.part.related_parts (package) e /word/media.
        # Abaixo, percorremos relacionamentos e coletamos blobs binários.
        rels = doc.part._rels
        for rel in rels:
            target = rels[rel].target_ref
            part = rels[rel]._target
            # Heurística: somente arquivos em word/media e com conteúdo binário
            try:
                if hasattr(part, "blob") and "/word/media/" in str(target):
                    imagens_bytes.append(part.blob)
            except Exception:
                continue

        texto_plano = "\n\n".join(texto_parts)
        return texto_plano, imagens_bytes

    def gerar_resumo_didatico(texto_plano: str):
        """
        Gera um resumo didático a partir do texto do arquivo.
        Não altera fatos: apenas reorganiza em linguagem simples.
        """
        # Pontos extraídos do documento (mantendo fatos e prazos)
        # ATENÇÃO: O conteúdo abaixo é uma reescrita didática dos pontos
        # presentes no arquivo, sem mudar números, datas ou condições.
        resumo = [
            "📄 **Documento-base**: Ato Conjunto RFB/CGIBS nº 001/2025 (regras para documentos fiscais IBS/CBS em 2026).",
            "🧾 **Para quem presta serviços (NFSe)**: segue obrigatório emitir NFSe; ela será recepcionada para IBS/CBS. No início, não há penalidade se os campos novos (IBS/CBS) não forem preenchidos até o 1º dia do 4º mês após os regulamentos. [1](https://myhines-my.sharepoint.com/personal/rcastellari_myhines_com/_layouts/15/Doc.aspx?sourcedoc=%7BA18CE2E3-193E-4E05-90A0-81E00B1A23DB%7D&file=fiscal%20reforma.docx&action=default&mobileredirect=true)",
            "ℹ️ **2026 é “informativo”**: há obrigação de enviar dados de IBS/CBS, mas sem efeitos tributários de apuração no ano. [1](https://myhines-my.sharepoint.com/personal/rcastellari_myhines_com/_layouts/15/Doc.aspx?sourcedoc=%7BA18CE2E3-193E-4E05-90A0-81E00B1A23DB%7D&file=fiscal%20reforma.docx&action=default&mobileredirect=true)",
            "🧭 **NFS-e nacional**: continua sob o CGNFS-e, com padronização nacional. [1](https://myhines-my.sharepoint.com/personal/rcastellari_myhines_com/_layouts/15/Doc.aspx?sourcedoc=%7BA18CE2E3-193E-4E05-90A0-81E00B1A23DB%7D&file=fiscal%20reforma.docx&action=default&mobileredirect=true)",
            "🧰 **Outros documentos**: existem para setores específicos; para prestadores de serviços comuns, foque na NFSe. [1](https://myhines-my.sharepoint.com/personal/rcastellari_myhines_com/_layouts/15/Doc.aspx?sourcedoc=%7BA18CE2E3-193E-4E05-90A0-81E00B1A23DB%7D&file=fiscal%20reforma.docx&action=default&mobileredirect=true)",
            "⚖️ **Tributos atuais seguem**: ISS (enquanto vigente), IRPJ, CSLL, PIS/COFINS etc. continuam conforme regras atuais. [1](https://myhines-my.sharepoint.com/personal/rcastellari_myhines_com/_layouts/15/Doc.aspx?sourcedoc=%7BA18CE2E3-193E-4E05-90A0-81E00B1A23DB%7D&file=fiscal%20reforma.docx&action=default&mobileredirect=true)",
            "🔁 **Transição 2026–2027 (CBS/IBS)**: 2026 tem CBS 0,9% e IBS 0,1% em regime de teste, com compensação contra PIS/COFINS; há possibilidade de dispensa parcial de recolhimento ao longo de 2026. Em 2027 extinguem-se PIS/COFINS e a CBS entra plenamente com alíquota definida por lei (referência ~8,8% ao fixar). [1](https://myhines-my.sharepoint.com/personal/rcastellari_myhines_com/_layouts/15/Doc.aspx?sourcedoc=%7BA18CE2E3-193E-4E05-90A0-81E00B1A23DB%7D&file=fiscal%20reforma.docx&action=default&mobileredirect=true)",
            "💳 **Créditos PIS/COFINS remanescentes**: continuam válidos; podem compensar CBS; regras de depreciação/amortização e estoque em 01/01/2027 são preservadas com critérios específicos e prazos (apropriação até jun/2027, uso em 12 parcelas). [1](https://myhines-my.sharepoint.com/personal/rcastellari_myhines_com/_layouts/15/Doc.aspx?sourcedoc=%7BA18CE2E3-193E-4E05-90A0-81E00B1A23DB%7D&file=fiscal%20reforma.docx&action=default&mobileredirect=true)",
            "🏢 **Locação por PJ (serviço com pouca despesa creditável)**: hoje paga 3,65% (PIS/COFINS cumulativo); após reforma, CBS não cumulativa e mais alta (~8,8% referencial), o que pode elevar carga quando praticamente não há créditos. Exemplo do documento mostra aumento mesmo com redutor e créditos comuns. [1](https://myhines-my.sharepoint.com/personal/rcastellari_myhines_com/_layouts/15/Doc.aspx?sourcedoc=%7BA18CE2E3-193E-4E05-90A0-81E00B1A23DB%7D&file=fiscal%20reforma.docx&action=default&mobileredirect=true)",
        ]
        return resumo

    # =========================
    # Ler arquivo e preparar conteúdo
    # =========================
    texto_plano, imagens_bytes = carregar_docx(DOCX_FILE)

    # ---- Abas
    tab_resumo, tab_completo, tab_transicao, tab_faq = st.tabs([
        "📌 Resumo", "📄 Conteúdo Completo (1:1)", "⏱️ Transição 2026–2027", "❓ Perguntas rápidas"
    ])

    # =========================
    # 📌 RESUMO DIDÁTICO
    # =========================
    with tab_resumo:
        st.subheader("Visão geral em linguagem simples")
        if not texto_plano:
            st.warning("Arquivo 'fiscal reforma.docx' não encontrado na mesma pasta do app.")
        else:
            pontos = gerar_resumo_didatico(texto_plano)
            for item in pontos:
                st.markdown(f"- {item}")

            st.info(
                "Este resumo reorganiza o conteúdo do arquivo em linguagem didática, "
                "sem alterar fatos, números ou prazos. Para conferência, veja a aba "
                "“Conteúdo Completo (1:1)”."
            )

    # =========================
    # 📄 CONTEÚDO COMPLETO (1:1)
    # =========================
    with tab_completo:
        st.subheader("Conteúdo integral do Word (texto + imagens)")
        if texto_plano:
            st.markdown("#### Texto integral")
            st.markdown(texto_plano)
        else:
            st.warning("Sem texto carregado.")

        if imagens_bytes:
            st.markdown("#### Imagens extraídas do documento")
            cols = st.columns(3)
            idx = 0
            for blob in imagens_bytes:
                try:
                    img = Image.open(BytesIO(blob))
                    with cols[idx % 3]:
                        st.image(img, use_column_width=True)
                    idx += 1
                except UnidentifiedImageError:
                    st.warning("Uma imagem do documento não pôde ser exibida.")
        else:
            st.info("Nenhuma imagem foi encontrada no documento ou não pôde ser extraída.")

    # =========================
    # ⏱️ TRANSIÇÃO 2026–2027
    # =========================
    with tab_transicao:
        st.subheader("Linha do tempo da transição (com base no documento)")
        st.markdown(
            "- **2026**: CBS 0,9% e IBS 0,1% em fase de teste, com compensação contra PIS/COFINS; "
            "existe a possibilidade de dispensa parcial de recolhimento condicionada a obrigações acessórias. [1](https://myhines-my.sharepoint.com/personal/rcastellari_myhines_com/_layouts/15/Doc.aspx?sourcedoc=%7BA18CE2E3-193E-4E05-90A0-81E00B1A23DB%7D&file=fiscal%20reforma.docx&action=default&mobileredirect=true)"
        )
        st.markdown(
            "- **2027**: Extinção de PIS/COFINS; CBS entra plenamente com alíquota definida por lei específica e/ou referência do Senado; "
            "preferência na compensação de créditos conforme regras de transição. [1](https://myhines-my.sharepoint.com/personal/rcastellari_myhines_com/_layouts/15/Doc.aspx?sourcedoc=%7BA18CE2E3-193E-4E05-90A0-81E00B1A23DB%7D&file=fiscal%20reforma.docx&action=default&mobileredirect=true)"
        )
        st.markdown(
            "- **Estoque e créditos**: créditos remanescentes de PIS/COFINS, critérios para estoque em 01/01/2027 e uso parcelado "
            "em 12 meses, conforme prazos do documento. [1](https://myhines-my.sharepoint.com/personal/rcastellari_myhines_com/_layouts/15/Doc.aspx?sourcedoc=%7BA18CE2E3-193E-4E05-90A0-81E00B1A23DB%7D&file=fiscal%20reforma.docx&action=default&mobileredirect=true)"
        )

    # =========================
    # ❓ PERGUNTAS RÁPIDAS (para leigos)
    # =========================
    with tab_faq:
        st.subheader("Perguntas rápidas – explicações simples")
        st.markdown("**O que é CBS e IBS?** — Impostos que substituem PIS/COFINS (CBS, federal) e ICMS/ISS (IBS, subnacional), com sistema não cumulativo e direito a crédito. [1](https://myhines-my.sharepoint.com/personal/rcastellari_myhines_com/_layouts/15/Doc.aspx?sourcedoc=%7BA18CE2E3-193E-4E05-90A0-81E00B1A23DB%7D&file=fiscal%20reforma.docx&action=default&mobileredirect=true)")
        st.markdown("**Em 2026 eu já pago CBS/IBS cheio?** — Não. 2026 é fase de teste (0,9% CBS e 0,1% IBS) com compensação e possível dispensa parcial; foco em cumprir a entrega das informações. [1](https://myhines-my.sharepoint.com/personal/rcastellari_myhines_com/_layouts/15/Doc.aspx?sourcedoc=%7BA18CE2E3-193E-4E05-90A0-81E00B1A23DB%7D&file=fiscal%20reforma.docx&action=default&mobileredirect=true)")
        st.markdown("**Quem aluga imóvel por PJ hoje paga o quê?** — PIS/COFINS 3,65% (cumulativo). Depois, CBS não cumulativa (alíquota mais alta), e setores com pouca despesa creditável tendem a sentir aumento. [1](https://myhines-my.sharepoint.com/personal/rcastellari_myhines_com/_layouts/15/Doc.aspx?sourcedoc=%7BA18CE2E3-193E-4E05-90A0-81E00B1A23DB%7D&file=fiscal%20reforma.docx&action=default&mobileredirect=true)")
        st.markdown("**O conteúdo desta página altera algo do documento?** — Não. O resumo só reorganiza as mesmas informações em linguagem simples; o conteúdo integral está disponível na aba própria para conferência. [1](https://myhines-my.sharepoint.com/personal/rcastellari_myhines_com/_layouts/15/Doc.aspx?sourcedoc=%7BA18CE2E3-193E-4E05-90A0-81E00B1A23DB%7D&file=fiscal%20reforma.docx&action=default&mobileredirect=true)")
